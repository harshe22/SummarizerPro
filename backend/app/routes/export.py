"""
Export API Routes
Handles document export and text-to-speech functionality
"""

from fastapi import APIRouter, HTTPException, Response
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, Dict, Any
import logging
import tempfile
import os
from pathlib import Path
import io

# Document generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

# Text-to-speech imports
from gtts import gTTS
import pyttsx3

logger = logging.getLogger(__name__)

router = APIRouter()

# Pydantic models
class ExportRequest(BaseModel):
    content: str
    title: Optional[str] = "Summary Report"
    format: str  # txt, docx, pdf
    metadata: Optional[Dict[str, Any]] = {}

class TTSRequest(BaseModel):
    text: str
    language: Optional[str] = "en"
    speed: Optional[float] = 1.0

@router.post("/export/txt")
async def export_txt(request: ExportRequest):
    """Export content as TXT file"""
    try:
        logger.info("Exporting content as TXT")
        
        if not request.content.strip():
            raise HTTPException(status_code=400, detail="Content cannot be empty")
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as temp_file:
            # Write header
            temp_file.write(f"{request.title}\n")
            temp_file.write("=" * len(request.title) + "\n\n")
            
            # Write metadata if available
            if request.metadata:
                temp_file.write("METADATA:\n")
                for key, value in request.metadata.items():
                    temp_file.write(f"{key.replace('_', ' ').title()}: {value}\n")
                temp_file.write("\n")
            
            # Write main content
            temp_file.write("CONTENT:\n")
            temp_file.write(request.content)
            
            temp_file_path = temp_file.name
        
        # Return file
        filename = f"{request.title.replace(' ', '_')}.txt"
        return FileResponse(
            path=temp_file_path,
            filename=filename,
            media_type='text/plain',
            background=lambda: os.unlink(temp_file_path)  # Clean up after sending
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error exporting TXT: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TXT export failed: {str(e)}")

@router.post("/export/docx")
async def export_docx(request: ExportRequest):
    """Export content as DOCX file"""
    try:
        logger.info("Exporting content as DOCX")
        
        if not request.content.strip():
            raise HTTPException(status_code=400, detail="Content cannot be empty")
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file_path = temp_file.name
        
        # Create DOCX document
        doc = Document()
        
        # Add title
        title = doc.add_heading(request.title, 0)
        title.alignment = 1  # Center alignment
        
        # Add metadata table if available
        if request.metadata:
            doc.add_heading('Summary Information', level=1)
            
            # Create metadata table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Property'
            hdr_cells[1].text = 'Value'
            
            # Add metadata rows
            for key, value in request.metadata.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key.replace('_', ' ').title()
                row_cells[1].text = str(value)
            
            doc.add_page_break()
        
        # Add main content
        doc.add_heading('Summary', level=1)
        
        # Split content into paragraphs
        paragraphs = request.content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save document
        doc.save(temp_file_path)
        
        # Return file
        filename = f"{request.title.replace(' ', '_')}.docx"
        return FileResponse(
            path=temp_file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            background=lambda: os.unlink(temp_file_path)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error exporting DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")

@router.post("/export/pdf")
async def export_pdf(request: ExportRequest):
    """Export content as PDF file"""
    try:
        logger.info("Exporting content as PDF")
        
        if not request.content.strip():
            raise HTTPException(status_code=400, detail="Content cannot be empty")
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file_path = temp_file.name
        
        # Create PDF document
        doc = SimpleDocTemplate(
            temp_file_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        # Build document content
        story = []
        
        # Add title
        story.append(Paragraph(request.title, title_style))
        story.append(Spacer(1, 12))
        
        # Add metadata table if available
        if request.metadata:
            story.append(Paragraph("Summary Information", heading_style))
            
            # Create metadata table
            table_data = [['Property', 'Value']]
            for key, value in request.metadata.items():
                table_data.append([
                    key.replace('_', ' ').title(),
                    str(value)
                ])
            
            table = Table(table_data, colWidths=[2.5*inch, 3.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            story.append(Spacer(1, 20))
        
        # Add main content
        story.append(Paragraph("Summary", heading_style))
        
        # Split content into paragraphs
        paragraphs = request.content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        
        # Return file
        filename = f"{request.title.replace(' ', '_')}.pdf"
        return FileResponse(
            path=temp_file_path,
            filename=filename,
            media_type='application/pdf',
            background=lambda: os.unlink(temp_file_path)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error exporting PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")

@router.post("/tts/generate")
async def generate_tts(request: TTSRequest):
    """Generate text-to-speech audio file"""
    try:
        logger.info("Generating TTS audio")
        
        if not request.text.strip():
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        
        if len(request.text) > 5000:
            raise HTTPException(status_code=400, detail="Text too long for TTS (max 5000 characters)")
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as temp_file:
            temp_file_path = temp_file.name
        
        try:
            # Use gTTS for text-to-speech
            tts = gTTS(
                text=request.text,
                lang=request.language,
                slow=False if request.speed >= 1.0 else True
            )
            
            # Save synchronously to avoid async issues
            tts.save(temp_file_path)
            
            # Verify file was created
            if not os.path.exists(temp_file_path) or os.path.getsize(temp_file_path) == 0:
                raise Exception("TTS file generation failed")
            
            # Define cleanup function
            def cleanup_file():
                try:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
                except Exception:
                    pass  # Ignore cleanup errors
            
            # Return audio file
            return FileResponse(
                path=temp_file_path,
                filename="summary_audio.mp3",
                media_type='audio/mpeg',
                background=cleanup_file
            )
            
        except Exception as gtts_error:
            logger.warning(f"gTTS failed: {str(gtts_error)}, trying pyttsx3")
            
            # Fallback to pyttsx3
            try:
                engine = pyttsx3.init()
                
                # Set properties
                rate = engine.getProperty('rate')
                engine.setProperty('rate', int(rate * request.speed))
                
                # Save to file
                engine.save_to_file(request.text, temp_file_path.replace('.mp3', '.wav'))
                engine.runAndWait()
                
                wav_path = temp_file_path.replace('.mp3', '.wav')
                
                # Define cleanup function for WAV file
                def cleanup_wav_file():
                    try:
                        if os.path.exists(wav_path):
                            os.unlink(wav_path)
                    except Exception:
                        pass  # Ignore cleanup errors
                
                return FileResponse(
                    path=wav_path,
                    filename="summary_audio.wav",
                    media_type='audio/wav',
                    background=cleanup_wav_file
                )
                
            except Exception as pyttsx3_error:
                logger.error(f"Both TTS engines failed: gTTS: {gtts_error}, pyttsx3: {pyttsx3_error}")
                raise HTTPException(status_code=500, detail="TTS generation failed with both engines")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating TTS: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TTS generation failed: {str(e)}")

@router.get("/tts/languages")
async def get_supported_languages():
    """Get list of supported TTS languages"""
    try:
        # Common gTTS supported languages
        supported_languages = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'it': 'Italian',
            'pt': 'Portuguese',
            'ru': 'Russian',
            'ja': 'Japanese',
            'ko': 'Korean',
            'zh': 'Chinese (Mandarin)',
            'ar': 'Arabic',
            'hi': 'Hindi',
            'nl': 'Dutch',
            'sv': 'Swedish',
            'no': 'Norwegian',
            'da': 'Danish',
            'fi': 'Finnish',
            'pl': 'Polish',
            'tr': 'Turkish',
            'th': 'Thai'
        }
        
        return {
            'supported_languages': supported_languages,
            'default_language': 'en',
            'total_languages': len(supported_languages)
        }
        
    except Exception as e:
        logger.error(f"Error getting supported languages: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to get supported languages")

@router.post("/export/batch")
async def batch_export(
    content: str,
    title: str = "Summary Report",
    formats: list = ["txt", "docx", "pdf"],
    metadata: dict = {}
):
    """Export content in multiple formats"""
    try:
        logger.info(f"Batch exporting content in formats: {formats}")
        
        if not content.strip():
            raise HTTPException(status_code=400, detail="Content cannot be empty")
        
        valid_formats = {"txt", "docx", "pdf"}
        invalid_formats = set(formats) - valid_formats
        
        if invalid_formats:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid formats: {invalid_formats}. Supported: {valid_formats}"
            )
        
        results = {}
        
        for format_type in formats:
            try:
                request = ExportRequest(
                    content=content,
                    title=title,
                    format=format_type,
                    metadata=metadata
                )
                
                if format_type == "txt":
                    result = await export_txt(request)
                elif format_type == "docx":
                    result = await export_docx(request)
                elif format_type == "pdf":
                    result = await export_pdf(request)
                
                results[format_type] = {
                    'status': 'success',
                    'filename': result.filename if hasattr(result, 'filename') else f"{title}.{format_type}"
                }
                
            except Exception as e:
                logger.error(f"Failed to export {format_type}: {str(e)}")
                results[format_type] = {
                    'status': 'error',
                    'error': str(e)
                }
        
        return {
            'results': results,
            'total_formats': len(formats),
            'successful_exports': len([r for r in results.values() if r['status'] == 'success']),
            'failed_exports': len([r for r in results.values() if r['status'] == 'error'])
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in batch export: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Batch export failed: {str(e)}")
